# ========================= docx_utils.py (Universal, auto-mapping) =========================
# TRF tick OK, SAMPLE PICTURE 3x4in (center), Summary & Detail auto-fill,
# Result style preserved. Supports ALL templates via TEMPLATE_MAP + TEST_GROUP_TITLES
# ===========================================================================================

import os
import re
import time
import uuid
import tempfile
import unicodedata
from io import BytesIO

from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from config import local_main, TEMPLATE_MAP
from test_logic import TEST_GROUP_TITLES
from excel_utils import _find_report_col

# Optional pandas dependency for cover fill from Excel
try:
    import pandas as pd
except Exception:
    pd = None

WORD_TEMPLATE = "FORM-QAD-011-TEST REQUEST FORM (TRF).docx"
PDF_OUTPUT_FOLDER = os.path.join("static", "TFR")
BLANK_TOKENS = {"", "-", "—", "–"}
_IMG_EXTS = (".png", ".jpg", ".jpeg", ".webp", ".bmp")

__all__ = [
    "get_first_empty_report_all_blank",
    "fill_docx_and_export_pdf",
    "approve_request_fill_docx_pdf",
    "fill_bed_cover_from_excel",
    "fill_cover_from_excel_generic",
    "create_report_for_type",
]

# ============================ Common helpers ============================

def _normalize_to_check_blank(v):
    if v is None:
        return True, ""
    if isinstance(v, str):
        s = (
            v.replace("\u00A0", "")
             .replace("\u200B", "")
             .replace("\r", "")
             .replace("\n", "")
             .replace("\t", "")
             .strip()
        )
        return (s in BLANK_TOKENS), s
    return False, str(v)

def _is_placeholder_dash(txt: str) -> bool:
    if txt is None:
        return True
    t = re.sub(r"[\s\r\n\t]+", "", str(txt or ""))
    return t in BLANK_TOKENS

def _is_result_placeholder(txt: str) -> bool:
    """Chỉ coi là placeholder nếu là '-' hoặc các dạng gạch ngang."""
    if txt is None:
        return True
    s = _norm(txt)
    return s in {"", "-", "—", "–"}

def _norm(s: str) -> str:
    s = unicodedata.normalize("NFD", str(s or ""))
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    s = s.lower()
    s = s.replace("không", "khong").replace("(không)", "(khong)").replace("(co)", "(có)")
    s = re.sub(r"[^a-z0-9]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _token_overlap(a: str, b: str) -> float:
    ta = set(_norm(a).split())
    tb = set(_norm(b).split())
    if not ta or not tb:
        return 0.0
    inter = len(ta & tb)
    return inter / max(1, min(len(ta), len(tb)))

def _clone_first_run_style(src_cell):
    try:
        p = src_cell.paragraphs[0]
        if not p.runs:
            return None, None, None, None
        r = p.runs[0]
        return r.font.name, r.font.size, r.font.bold, r.font.italic
    except Exception:
        return None, None, None, None

def _apply_text_with_font(dst_cell, text, fname=None, fsize=None, fbold=None, fitalic=None, align=None, extra_bottom_text=None):
    for p in list(dst_cell.paragraphs):
        if hasattr(p, "clear"):
            try:
                p.clear()
            except Exception:
                pass
    dst_cell.text = ""

    p1 = dst_cell.paragraphs[0] if dst_cell.paragraphs else dst_cell.add_paragraph("")
    run1 = p1.add_run("" if text is None else str(text))
    if fname:   run1.font.name = fname
    if fsize:   run1.font.size = fsize
    if fbold is not None:   run1.font.bold = fbold
    if fitalic is not None: run1.font.italic = fitalic
    if align is not None:   p1.alignment = align
    pf = p1.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)

    if extra_bottom_text:
        p2 = dst_cell.add_paragraph("")
        run2 = p2.add_run(str(extra_bottom_text))
        if fname:   run2.font.name = fname
        if fsize:   run2.font.size = fsize
        run2.font.bold = False  # comment không in đậm
        if fitalic is not None: run2.font.italic = fitalic
        if align is not None:   p2.alignment = align
        pf2 = p2.paragraph_format
        pf2.space_before = Pt(0)
        pf2.space_after  = Pt(0)

def _set_cell_text_with_style(dst_cell, src_label_cell, text, align_center=False):
    fname, fsize, fbold, fitalic = _clone_first_run_style(src_label_cell)
    _apply_text_with_font(
        dst_cell,
        text,
        fname, fsize, fbold, fitalic,
        align=WD_ALIGN_PARAGRAPH.CENTER if align_center else None
    )

# ============================ Excel (TRF) ============================

def get_first_empty_report_all_blank(excel_path):
    wb = load_workbook(excel_path, data_only=True)
    ws = wb.active
    report_col = _find_report_col(ws)

    for row in range(2, ws.max_row + 1):
        all_mid_empty = True
        for col in range(3, 25):  # C..X
            is_blank, _ = _normalize_to_check_blank(ws.cell(row=row, column=col).value)
            if not is_blank:
                all_mid_empty = False
                break
        if all_mid_empty:
            report_no = ws.cell(row=row, column=report_col).value
            if report_no is not None and str(report_no).strip():
                wb.close()
                return str(report_no).strip()
    wb.close()
    return None

# ============================ Unicode checkboxes ============================

def _label_regex(label: str) -> re.Pattern:
    cleaned = re.sub(r'[_\.\-]+', ' ', (label or '').strip())
    parts = [p for p in cleaned.split() if p]
    pattern = r'(☐|☑)\s*' + r'\s*'.join(re.escape(p) for p in parts)
    return re.compile(pattern, flags=re.IGNORECASE)

def build_label_value_map(data):
    label_groups = {
        "test_group": [
            "MATERIAL TEST",
            "FINISHING TEST",
            "CONSTRUCTION TEST",
            "TRANSIT TEST",
            "ENVIRONMENTAL TEST",
        ],
        "test_status": ["1ST", "2ND", "3RD", "...TH"],
        "furniture_testing": ["INDOOR", "OUTDOOR"],
        "sample_return": ["YES", "NO"],
    }

    def _eq_relaxed(label: str, value: str, group: str) -> bool:
        L = (label or "").strip().upper()
        V = (value or "").strip().upper()
        if not V:
            return False
        if group == "test_group":
            V2 = V[:-5].strip() if V.endswith(" TEST") else V
            L2 = L[:-5].strip() if L.endswith(" TEST") else L
            return (V == L) or (V2 == L) or (V == L2) or (V2 == L2)
        if group == "test_status" and label == "...TH":
            return V.endswith("TH") and V not in {"1ST", "2ND", "3RD"}
        return V == L

    label_value_map = {}
    for group, labels in label_groups.items():
        value = data.get(group, None)
        if value is None or (isinstance(value, str) and not value.strip()):
            for label in labels:
                label_value_map[label] = False
        else:
            for label in labels:
                label_value_map[label] = _eq_relaxed(label, str(value), group)

    for field in ["sample_description", "item_code", "supplier", "subcon"]:
        val = str(data.get(field, "")).strip().upper()
        label_value_map[f"{field.upper()} N/A"] = (val == "N/A")

    return label_value_map

def tick_unicode_checkbox_by_label(doc: Document, label_value_map):
    compiled = [(_label_regex(label), bool(value)) for label, value in label_value_map.items()]

    def toggle_text(txt: str) -> str:
        if not txt or ('☐' not in txt and '☑' not in txt):
            return txt
        for pat, value in compiled:
            def _repl(m):
                return ('☑' if value else '☐') + m.group(0)[1:]
            txt = pat.sub(_repl, txt)
        return txt

    for p in doc.paragraphs:
        t = p.text
        if ("☐" in t) or ("☑" in t):
            p.text = toggle_text(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text
                if ("☐" in t) or ("☑" in t):
                    cell.text = toggle_text(t)

# ============================ PDF convert (optional, Windows) ============================

def try_convert_to_pdf(docx_path, pdf_path):
    try:
        import pythoncom
        pythoncom.CoInitialize()
        from docx2pdf import convert
        convert(docx_path, pdf_path)
    except Exception as e:
        import traceback
        print("PDF convert failed:", e)
        traceback.print_exc()

# ============================ Atomic save ============================

def _lock_path_for(report_no: str) -> str:
    return os.path.join(tempfile.gettempdir(), f"tfr_{report_no}.lock")

def _acquire_lock(path: str, timeout=30):
    t0 = time.time()
    fd = None
    while True:
        try:
            fd = os.open(path, os.O_CREAT | os.O_EXCL | os.O_RDWR)
            return fd
        except FileExistsError:
            if time.time() - t0 > timeout:
                try:
                    os.unlink(path)
                except:
                    pass
            time.sleep(0.05)
        except Exception:
            time.sleep(0.05)

def _release_lock(fd, path: str):
    try:
        os.close(fd)
    except:
        pass
    try:
        os.unlink(path)
    except:
        pass

def _atomic_save_docx(doc: Document, out_path: str):
    tmp = f"{out_path}.tmp-{uuid.uuid4().hex}"
    doc.save(tmp)
    os.replace(tmp, out_path)

# ============================ TRF: fill & export ============================

def fill_docx_and_export_pdf(data, fixed_report_no=None):
    # report no
    if fixed_report_no and str(fixed_report_no).strip():
        report_no = str(fixed_report_no).strip()
    else:
        report_no = get_first_empty_report_all_blank(_smart_excel_path(local_main))
        if not report_no:
            raise Exception("No empty report number available in Excel.")

    data = dict(data or {})
    data["report_no"] = report_no
    template_key = (data.get("template_key") or "other")

    doc = Document(WORD_TEMPLATE)

    # basic field mapping
    mapping = {
        "requestor": "requestor",
        "department": "department",
        "requested date": "request_date",
        "lab test report no.": "report_no",
        "sample description": "sample_description",
        "item code": "item_code",
        "quantity": "quantity",
        "supplier": "supplier",
        "subcon": "subcon",
        "test group": "test_group",
        "test status": "test_status",
        "furniture testing": "furniture_testing",
        "estimated completed date": "estimated_completion_date",
    }
    remark = data.get("remark", "")
    remark_written = False

    for table in doc.tables:
        nrows = len(table.rows)
        ncols = len(table.columns)
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                label = (
                    cell.text.strip().lower()
                    .replace("(mã item)", "")
                    .replace("(mã material)", "")
                    .replace("*", "")
                )
                if not remark_written and ("other tests/instructions" in label or "remark" in label) and remark:
                    if i + 1 < nrows:
                        below_cell = table.rows[i + 1].cells[j]
                        if not (below_cell.text or "").strip():
                            below_cell.text = str(remark)
                            remark_written = True
                            continue
                if ("emp id" in label or "msnv" in label) and data.get("employee_id", ""):
                    if j + 1 < ncols:
                        target_cell = row.cells[j + 1]
                        if not (target_cell.text or "").strip():
                            target_cell.text = str(data["employee_id"])
                            continue
                for map_label, key in mapping.items():
                    if map_label in ["remark", "employee id"]:
                        continue
                    if map_label in label and key in data and str(data[key]).strip() != "":
                        if j + 1 < ncols:
                            target_cell = row.cells[j + 1]
                            if (target_cell.text or "").strip() == "" or "lab test report no." in label:
                                target_cell.text = str(data[key])

    # Tick boxes
    label_value_map = build_label_value_map(data)
    tick_unicode_checkbox_by_label(doc, label_value_map)

    # === NEW: auto fill Summary/Detail (status + comment + photo) for TRF too ===
    _update_exec_summary_results_from_status(doc, report_no, template_key)
    _update_detail_results_and_comments(doc, report_no, template_key)

    # save
    if not os.path.exists(PDF_OUTPUT_FOLDER):
        os.makedirs(PDF_OUTPUT_FOLDER)
    output_docx = os.path.join(PDF_OUTPUT_FOLDER, f"{report_no}.docx")
    output_pdf  = os.path.join(PDF_OUTPUT_FOLDER, f"{report_no}.pdf")

    lock_path = _lock_path_for(report_no)
    fd = _acquire_lock(lock_path, timeout=30)
    try:
        _atomic_save_docx(doc, output_docx)
    finally:
        _release_lock(fd, lock_path)

    try_convert_to_pdf(output_docx, output_pdf)
    return output_docx, output_pdf, report_no


def approve_request_fill_docx_pdf(req):
    fixed = (req.get("report_no") or "").strip()
    if "etd" in req and not req.get("estimated_completion_date"):
        req = dict(req)
        req["estimated_completion_date"] = req.get("etd")

    if fixed:
        out_docx, out_pdf, report_no = fill_docx_and_export_pdf(req, fixed_report_no=fixed)
    else:
        out_docx, out_pdf, report_no = fill_docx_and_export_pdf(req, fixed_report_no=None)

    if os.path.exists(out_pdf):
        return out_pdf, report_no
    return out_docx, report_no

# ======================= SAMPLE PICTURE =======================

def _find_overview_images(report_id: str) -> list[str]:
    roots = []
    images_root = os.path.join(os.getcwd(), "images")
    report_dods_root = os.path.join(os.getcwd(), "report dods")
    if report_id:
        roots.append(os.path.join(images_root, str(report_id)))
        roots.append(os.path.join(report_dods_root, str(report_id)))
    roots.append(report_dods_root)
    roots.append(images_root)

    found = []
    seen = set()
    for root in roots:
        if not os.path.isdir(root):
            continue
        for name in os.listdir(root):
            low = name.lower()
            if any(low.endswith(ext) for ext in _IMG_EXTS) and "overview" in low:
                p = os.path.join(root, name)
                if p not in seen:
                    seen.add(p)
                    found.append(p)
        if found:
            break

    found.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    return found

def _find_sample_picture_target_cell(doc: Document):
    for t in doc.tables:
        rows = t.rows
        for i, row in enumerate(rows):
            for j, cell in enumerate(row.cells):
                if "sample picture" in _norm(cell.text):
                    if i + 1 < len(rows):
                        return rows[i + 1].cells[0]
    return None

def _clear_cell_keep_one_paragraph(cell):
    try:
        cell.text = ""
        if not cell.paragraphs:
            cell.add_paragraph("")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

def _insert_overview_images_into_sample_picture(doc: Document, report_id: str) -> bool:
    img_paths = _find_overview_images(report_id)
    if not img_paths:
        return False

    target_cell = _find_sample_picture_target_cell(doc)
    if target_cell is None:
        return False

    _clear_cell_keep_one_paragraph(target_cell)
    target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    pic_w = Inches(3.5)
    pic_h = Inches(2.5)

    sect = doc.sections[0]
    avail_w = int((sect.page_width - sect.left_margin - sect.right_margin) * 0.97)
    cols = max(1, int(avail_w // pic_w))

    # tạo table con
    inner = target_cell.add_table(rows=0, cols=cols)
    inner.alignment = WD_ALIGN_PARAGRAPH.CENTER   # ép cả table con căn giữa
    inner.autofit = False

    for c in inner.columns:
        for cell in c.cells:
            cell.width = pic_w
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    r = None
    for idx, path in enumerate(img_paths):
        if idx % cols == 0:
            r = inner.add_row()
        c = r.cells[idx % cols]
        c.text = ""
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0] if c.paragraphs else c.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER   # căn giữa ảnh trong cell
        run = p.add_run()
        pic = run.add_picture(path)
        pic.width = pic_w
        pic.height = pic_h

    return True

# ======================= STATUS & COMMENT & PHOTO =======================

def _find_status_file(report_id: str) -> str | None:
    def _candidates(root_dir):
        if not os.path.isdir(root_dir):
            return []
        return [
            os.path.join(root_dir, f)
            for f in os.listdir(root_dir)
            if f.lower().startswith("status") and f.lower().endswith(".txt")
        ]

    base_roots = [os.path.join(os.getcwd(), "images"), os.path.join(os.getcwd(), "report dods")]
    pri_roots = []
    if report_id:
        for b in base_roots:
            pri_roots.append(os.path.join(b, str(report_id)))

    cand = []
    for r in (pri_roots + base_roots):
        if os.path.isdir(r):
            cand.extend(_candidates(r))
    if not cand:
        return None
    cand.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    return cand[0]

def _find_comment_file(report_id: str) -> str | None:
    def _candidates(root_dir):
        if not os.path.isdir(root_dir):
            return []
        return [
            os.path.join(root_dir, f)
            for f in os.listdir(root_dir)
            if f.lower().startswith("comment") and f.lower().endswith(".txt")
        ]

    base_roots = [os.path.join(os.getcwd(), "images"), os.path.join(os.getcwd(), "report dods")]
    pri_roots = []
    if report_id:
        for b in base_roots:
            pri_roots.append(os.path.join(b, str(report_id)))

    cand = []
    for r in (pri_roots + base_roots):
        if os.path.isdir(r):
            cand.extend(_candidates(r))
    if not cand:
        return None
    cand.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    return cand[0]

def _read_status_map(report_id: str) -> dict:
    fp = _find_status_file(report_id)
    out = {}
    if not fp or not os.path.exists(fp):
        return out
    with open(fp, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            # chấp nhận cả muc4.2, muc8.3.3
            m = re.match(r"^\s*(muc[\d\.]+)\s*:\s*([A-Za-z/ ]+)\s*$", line.strip())
            if not m:
                continue
            key = m.group(1).lower()
            val = m.group(2).strip().upper()
            if val == "NA":
                val = "N/A"
            out[key] = val
    return out


def _read_comment_map(report_id: str) -> dict:
    fp = _find_comment_file(report_id)
    out = {}
    if not fp or not os.path.exists(fp):
        return out
    with open(fp, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            # chấp nhận cả muc4.2, muc8.3.3
            m = re.match(r"^\s*(muc[\d\.]+)\s*:\s*(.+?)\s*$", line.strip())
            if not m:
                continue
            key = m.group(1).lower()
            val = m.group(2).strip()
            out[key] = val
    return out

# --------- Resolve TEST_GROUP_TITLES key from template_key (auto-mapping) ---------

def _resolve_group_key(template_key: str) -> str | None:
    if not template_key:
        return None
    available_keys = list(TEST_GROUP_TITLES.keys())
    lc_to_orig = {k.lower(): k for k in available_keys}
    direct = lc_to_orig.get(template_key.lower())
    if direct:
        return direct

    alias_map = {
        "bed": "giuong",
        "chair_us": "ghe_us",
        "chair_uk": "ghe_eu",
        "table_us": "ban_us",
        "table_uk": "ban_eu",
        "cabinet_us": "tu_us",
        "cabinet_uk": "tu_eu",
        "mirror": "guong",
        "material_indoor_chuyen": "indoor_chuyen",
        "material_indoor_qa": "indoor_thuong",
        "material_indoor_stone": "indoor_stone",
        "material_indoor_metal": "indoor_metal",
        "material_outdoor": "outdoor_finishing",
        "line_test": "line",
        "transit_2c_np": "transit_2c_np",
        "transit_rh_np": "transit_rh_np",
        "transit_181_lt68": "transit_181_lt68",
        "transit_3a": "transit_3a",
        "transit_3b_np": "transit_3b_np",
        "transit_2c_pallet": "transit_2c_pallet",
        "transit_rh_pallet": "transit_rh_pallet",
        "transit_181_gt68": "transit_181_gt68",
        "transit_3b_pallet": "transit_3b_pallet",
        "hot_cold_test": "hot_cold_test",
        "other": "other",
    }
    alias_target = alias_map.get(template_key.lower())
    if alias_target:
        return lc_to_orig.get(alias_target.lower(), alias_target)

    def _score_key(k):
        return _token_overlap(template_key, k) + _token_overlap(template_key.replace("uk", "eu"), k)

    best_key, best_score = None, 0.0
    for k in available_keys:
        sc = _score_key(k)
        if sc > best_score:
            best_key, best_score = k, sc
    return best_key if best_score >= 0.45 else None

def _prep_title_candidates(template_key: str) -> dict:
    resolved = _resolve_group_key(template_key)
    titles_map = TEST_GROUP_TITLES.get(resolved or "", {}) or {}

    cand = {}
    for muc, info in titles_map.items():
        names = set()
        for k in ("full", "short"):
            v = info.get(k)
            if v:
                names.add(v)
        more = set()
        for n in list(names):
            x = n
            x = re.sub(r"\btest\b", "", x, flags=re.IGNORECASE)
            x = re.sub(r"[\(\)]", " ", x)
            more.add(x)
        names |= more
        cand[muc.lower()] = {_norm(n) for n in names if n}
    return cand

def _match_muc(desc: str, clause: str, muc_cands: dict) -> str | None:
    parts = [desc or ""]
    if clause:
        parts.append(clause)
    nd_full = _norm(" ".join(parts))

    best_muc, best_score = None, 0.0
    for muc, patterns in muc_cands.items():
        score = 0.0
        for p in patterns:
            if not p:
                continue
            if p in nd_full or nd_full in p:
                score = max(score, 1.0)
            else:
                score = max(score, _token_overlap(nd_full, p))
        if score > best_score:
            best_muc, best_score = muc, score
    return best_muc if best_score >= 0.4 else None  # relaxed threshold

# ======================= Table detection (with synonyms) =======================

def _find_exec_summary_table(doc: Document):
    """
    Nhận diện bảng EXECUTIVE SUMMARY với các alias linh hoạt:
      - Cột 1: "Clause" hoặc "Test property"
      - Cột 2: "Description"
      - Cột 3: "Result"
      - Cột 4 (tuỳ chọn): "Comment(s)" / "*Comments"
    """
    clause_aliases = {"clause", "test property"}
    desc_aliases = {"description"}
    result_aliases = {"result"}
    comment_aliases = {"comment", "comments", "*comments"}

    for t in doc.tables:
        for r in t.rows:
            cells = r.cells
            if len(cells) >= 3:
                heads = [_norm(cells[i].text) for i in range(min(len(cells), 4))]
                if (
                    len(heads) > 2
                    and heads[0] in clause_aliases
                    and heads[1] in desc_aliases
                    and heads[2] in result_aliases
                ):
                    return t
    return None


def _find_detail_table(doc: Document):
    """
    Nhận diện bảng chi tiết:
      - Cột 1: "Clause" hoặc "Test property"
      - Cột 2: "Description"
      - Cột 3: "Test Method/Requirement" hoặc "Criteria"
      - Cột 4: "Result"
      - Cột 5 (tuỳ chọn): "Photo reference"
    """
    clause_aliases = {"clause", "test property"}
    desc_aliases = {"description"}
    req_aliases = {"test method requirement", "criteria", "test method", "requirement"}
    result_aliases = {"result"}
    photo_aliases = {"photo reference", "photo"}

    for t in doc.tables:
        for r in t.rows:
            cells = r.cells
            if len(cells) >= 4:
                heads = [_norm(cells[i].text) for i in range(min(len(cells), 5))]
                if (
                    len(heads) > 3
                    and heads[0] in clause_aliases
                    and heads[1] in desc_aliases
                    and heads[2] in req_aliases
                    and heads[3] in result_aliases
                ):
                    return t
    return None

# ======================= Result/Photo style detection =======================

def _detect_result_style(tbl):
    for r in tbl.rows:
        cells = r.cells
        if len(cells) < 3:
            continue
        c = cells[2]
        t = (c.text or "").strip()
        if t and not _is_placeholder_dash(t):
            fname, fsize, fbold, fitalic = _clone_first_run_style(c)
            try:
                align = c.paragraphs[0].alignment or WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                align = WD_ALIGN_PARAGRAPH.CENTER
            return fname, fsize, fbold, fitalic, align
    return None, None, None, None, WD_ALIGN_PARAGRAPH.CENTER

def _detect_result_style_detail(tbl):
    for r in tbl.rows:
        cells = r.cells
        if len(cells) < 4:
            continue
        c = cells[3]
        t = (c.text or "").strip()
        if t and not _is_placeholder_dash(t):
            fname, fsize, fbold, fitalic = _clone_first_run_style(c)
            try:
                align = c.paragraphs[0].alignment or WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                align = WD_ALIGN_PARAGRAPH.CENTER
            return fname, fsize, fbold, fitalic, align
    return None, None, None, None, WD_ALIGN_PARAGRAPH.CENTER

def _detect_photo_style_detail(tbl):
    for r in tbl.rows:
        cells = r.cells
        if len(cells) < 5:
            continue
        c = cells[4]
        t = (c.text or "").strip()
        if t and t.upper() not in {"-", "NO PHOTO"}:
            fname, fsize, fbold, fitalic = _clone_first_run_style(c)
            try:
                align = c.paragraphs[0].alignment or WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                align = WD_ALIGN_PARAGRAPH.CENTER
            return fname, fsize, fbold, fitalic, align
        if (t or "").upper() == "NO PHOTO":
            try:
                align = c.paragraphs[0].alignment or WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                align = WD_ALIGN_PARAGRAPH.CENTER
            return None, None, None, None, align
    return None, None, None, None, WD_ALIGN_PARAGRAPH.CENTER

# ======================= Fill Summary & Detail =======================

def _update_exec_summary_results_from_status(doc: Document, report_id: str, template_key: str) -> bool:
    status_map = _read_status_map(report_id)
    if not status_map:
        return False

    muc_cands = _prep_title_candidates(template_key)
    if not muc_cands:
        return False

    tbl = _find_exec_summary_table(doc)
    if tbl is None:
        return False

    fname, fsize, fbold, fitalic, align = _detect_result_style(tbl)

    changed = False
    for i, r in enumerate(tbl.rows):
        if i == 0:  # luôn bỏ header
            continue
        cells = r.cells
        if len(cells) < 3:
            continue

        clause_text = cells[0].text or ""
        desc_text = cells[1].text or ""
        result_cell = cells[2]

        if not _is_result_placeholder(result_cell.text):
            continue

        muc = _match_muc(desc_text, clause_text, muc_cands)
        if not muc:
            continue

        val = status_map.get(muc)
        if not val:
            continue

        _apply_text_with_font(result_cell, val, fname, fsize, fbold, fitalic, align=align)
        changed = True
    return changed

# --------- Image helpers ---------

def _all_candidate_images(report_id: str):
    roots = []
    images_root = os.path.join(os.getcwd(), "images")
    report_dods_root = os.path.join(os.getcwd(), "report dods")
    if report_id:
        roots.append(os.path.join(images_root, str(report_id)))
        roots.append(os.path.join(report_dods_root, str(report_id)))
    roots.append(report_dods_root)
    roots.append(images_root)

    files = []
    seen = set()
    for root in roots:
        if not os.path.isdir(root):
            continue
        for name in os.listdir(root):
            low = name.lower()
            if any(low.endswith(ext) for ext in _IMG_EXTS):
                p = os.path.join(root, name)
                if p not in seen:
                    seen.add(p)
                    files.append(p)
    files.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    return files

def _extract_muc_and_order_from_name(path: str, known_mucs: set[str] | None = None) -> tuple[str | None, int]:
    """
    Parse tên file thành (muc_key, order).
    Ví dụ:
      - ..._muc4.3_1.png  → ('muc4.3', 1)
      - ..._muc8.2.2_3.jpg → ('muc8.2.2', 3)
      - ..._muc5.7.1.jpeg  → nếu 'muc5.7' ∈ known_mucs và 'muc5.7.1' ∉ known_mucs
                              thì hiểu là ('muc5.7', 1)
    """
    base = os.path.basename(path).lower()
    name_noext = re.sub(r"\.[a-z0-9]+$", "", base)

    # Case A: mucX[.Y...][_ - ]order
    m = re.search(r"muc(?P<muc>\d+(?:\.\d+)*)(?:[ _\-]+(?P<ord>\d+))\b", name_noext)
    if m:
        muc = f"muc{m.group('muc')}"
        order = int(m.group('ord'))
        return muc, order

    # Case B: mucA.B.order (không có underscore)
    m2 = re.search(r"muc(?P<parts>\d+(?:\.\d+)+)\b", name_noext)
    if m2:
        parts = m2.group("parts").split(".")
        full_key = "muc" + ".".join(parts)
        muc, order = full_key, 1
        if len(parts) >= 2:
            maybe_parent = "muc" + ".".join(parts[:-1])
            maybe_ord = parts[-1]
            if maybe_ord.isdigit():
                if known_mucs and (maybe_parent in known_mucs) and (full_key not in known_mucs):
                    muc = maybe_parent
                    order = int(maybe_ord)
        return muc, order

    return None, 0

_IMAGE_INDEX_CACHE: dict[tuple[str, tuple[str, ...]], dict[str, list[str]]] = {}

def _index_images_by_muc(report_id: str, known_mucs: set[str]) -> dict[str, list[str]]:
    """
    Lập chỉ mục {muc → [ảnh1, ảnh2,...]} (ảnh đã sort theo order).
    """
    cache_key = (str(report_id), tuple(sorted(known_mucs)))
    if cache_key in _IMAGE_INDEX_CACHE:
        return _IMAGE_INDEX_CACHE[cache_key]

    paths = _all_candidate_images(report_id)
    buckets: dict[str, list[tuple[int, str]]] = {}

    for p in paths:
        muc, order = _extract_muc_and_order_from_name(p, known_mucs)
        if not muc:
            continue
        if muc not in buckets:
            buckets[muc] = []
        buckets[muc].append((order if order > 0 else 1, p))

    out: dict[str, list[str]] = {}
    for muc, items in buckets.items():
        items.sort(key=lambda it: (it[0], os.path.getmtime(it[1])))
        out[muc] = [path for _, path in items]

    _IMAGE_INDEX_CACHE[cache_key] = out
    return out

def _insert_photo_references_stack(cell, image_paths: list[str], fname=None, fsize=None, fbold=None, fitalic=None, align=None):
    for p in list(cell.paragraphs):
        if hasattr(p, "clear"):
            try:
                p.clear()
            except Exception:
                pass
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if not image_paths:
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph("")
        run = p.add_run("NO PHOTO")
        run.font.bold = False
        if fname:   run.font.name = fname
        if fsize:   run.font.size = fsize
        if fitalic is not None: run.font.italic = fitalic
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align is None else align
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        return

    for idx, path in enumerate(image_paths):
        p = cell.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align is None else align
        run = p.add_run()
        try:
            pic = run.add_picture(path)
            pic.width = Inches(2.0)
            pic.height = Inches(2.5)
        except Exception:
            continue
        pf = p.paragraph_format
        pf.space_before = Pt(0 if idx == 0 else 4)
        pf.space_after  = Pt(0)

def _update_detail_results_and_comments(doc: Document, report_id: str, template_key: str) -> bool:
    status_map = _read_status_map(report_id)
    if not status_map:
        return False
    comment_map = _read_comment_map(report_id)

    muc_cands = _prep_title_candidates(template_key)
    if not muc_cands:
        return False

    tbl = _find_detail_table(doc)
    if tbl is None:
        return False

    known_mucs = set(status_map.keys()) | set(comment_map.keys())
    img_index = _index_images_by_muc(report_id, known_mucs)

    fname_r, fsize_r, fbold_r, fitalic_r, align_r = _detect_result_style_detail(tbl)
    fname_p, fsize_p, fbold_p, fitalic_p, align_p = _detect_photo_style_detail(tbl)

    changed = False
    for i, r in enumerate(tbl.rows):
        if i == 0:
            continue
        cells = r.cells
        if len(cells) < 5:
            continue

        clause_text = cells[0].text or ""
        desc_text = cells[1].text or ""
        result_cell = cells[3]
        photo_cell  = cells[4]

        muc = _match_muc(desc_text, clause_text, muc_cands)

        # Result
        if muc:
            cur_val = (result_cell.text or "").strip()
            if _is_result_placeholder(cur_val):
                val = status_map.get(muc)
                if val:
                    extra = comment_map.get(muc, None)
                    _apply_text_with_font(
                        result_cell, val,
                        fname_r, fsize_r, fbold_r, fitalic_r,
                        align=align_r, extra_bottom_text=extra
                    )
                    changed = True

        # Photo
        cur_text = (photo_cell.text or "").strip().upper()
        if cur_text != "NO PHOTO" and (_is_result_placeholder(cur_text) or cur_text == ""):
            photos = img_index.get(muc or "", []) if muc else []
            _insert_photo_references_stack(photo_cell, photos, fname_p, fsize_p, fbold_p, fitalic_p, align_p)
            changed = True

        if not muc:
            cur_text2 = (photo_cell.text or "").strip()
            if _is_result_placeholder(cur_text2) or cur_text2 == "":
                _insert_photo_references_stack(photo_cell, [], fname_p, fsize_p, fbold_p, fitalic_p, align_p)
                changed = True

    return changed

# ======================= Cover table (from Excel) =======================

def _find_result_table(doc: Document):
    cover_labels = {
        "Sample Description:", "Item/ Material code:", "Category:", "Collection:",
        "Country of Destination:", "Supplier/ Subcontractor:", "Customer:",
        "Sample Size:", "Sample Weight:", "Tested by:", "Generated by:"
    }
    for t in doc.tables:
        if len(t.columns) >= 4:
            seen = set()
            for r in t.rows:
                for c in r.cells:
                    txt = (c.text or "").strip()
                    if txt in cover_labels:
                        seen.add(txt)
            if len(seen) >= 4:
                return t
    return doc.tables[0] if doc.tables else None

def _set_result_inline_in_paragraph(paragraph, rating_text: str) -> bool:
    if not rating_text or not paragraph.runs:
        return False

    full = "".join(r.text for r in paragraph.runs)
    m = re.search(r"(?i)RESULT\s*:\s*", full)
    if not m:
        return False
    start_after_colon = m.end()

    idx = 0
    run_index = None
    offset_in_run = None
    for ridx, r in enumerate(paragraph.runs):
        txt = r.text
        nxt = idx + len(txt)
        if nxt > start_after_colon and idx <= start_after_colon:
            offset_in_run = start_after_colon - idx
            tail = txt[offset_in_run:]
            m2 = re.match(r"[\s]*[–—-]+", tail)
            if m2:
                run_index = ridx
            break
        idx = nxt

    if run_index is None:
        for ridx in range(len(paragraph.runs)):
            if re.fullmatch(r"\s*[–—-]+\s*", paragraph.runs[ridx].text or ""):
                run_index = ridx
                offset_in_run = 0
                break

    if run_index is None:
        return False

    r = paragraph.runs[run_index]
    txt = r.text or ""
    head = txt[:offset_in_run] if offset_in_run else ""
    new_txt = re.sub(
        r"^\s*[–—-]+\s*",
        " " + str(rating_text),
        (txt[offset_in_run:] if offset_in_run is not None else txt),
    )
    r.text = head + new_txt
    return True

def _set_result_value(doc: Document, rating_text: str):
    if not rating_text:
        return False

    for t in doc.tables:
        for r in t.rows:
            cells = r.cells
            for j in range(len(cells) - 1):
                if (cells[j].text or "").strip().upper().startswith("RESULT:"):
                    cur = (cells[j + 1].text or "").strip()
                    if cur in BLANK_TOKENS or _is_result_placeholder(cur):
                        _set_cell_text_with_style(cells[j + 1], cells[j], rating_text)
                        return True

    for p in doc.paragraphs:
        if "RESULT" in (p.text or "") or "Result" in (p.text or ""):
            if _set_result_inline_in_paragraph(p, rating_text):
                return True
    return False

# ======================= Excel path resolution =======================

def _smart_excel_path(path_or_name: str) -> str:
    def _is_excel_file(p):
        low = p.lower()
        return os.path.isfile(p) and (low.endswith(".xlsx") or low.endswith(".xls"))

    def _latest_excel_in_dir(d):
        cands = []
        try:
            for name in os.listdir(d):
                fp = os.path.join(d, name)
                if _is_excel_file(fp):
                    cands.append(fp)
        except Exception:
            pass
        if not cands:
            return ""
        cands.sort(key=lambda p: os.path.getmtime(p), reverse=True)
        return cands[0]

    if path_or_name:
        p = str(path_or_name)
        if _is_excel_file(p):
            return p
        if os.path.isdir(p):
            latest = _latest_excel_in_dir(p)
            if latest:
                return latest
        if os.path.isdir(local_main):
            cand = os.path.join(local_main, p)
            if _is_excel_file(cand):
                return cand
            latest = _latest_excel_in_dir(local_main)
            if latest:
                return latest
        if _is_excel_file(local_main):
            return local_main
        raise FileNotFoundError(f"Excel not found: {path_or_name}")
    else:
        if os.path.isdir(local_main):
            latest = _latest_excel_in_dir(local_main)
            if latest:
                return latest
            raise FileNotFoundError(f"No Excel files found in directory: {local_main}")
        if _is_excel_file(local_main):
            return local_main
        raise FileNotFoundError("Excel path is not provided and local_main is not a valid Excel file or directory.")

def _load_excel_df(excel_path_or_name: str):
    if pd is None:
        raise RuntimeError("pandas is not available")
    excel_path = _smart_excel_path(excel_path_or_name)
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"Excel not found: {excel_path}")
    df = pd.read_excel(excel_path, sheet_name=0)
    df.columns = [str(c).strip() for c in df.columns]
    return df

# ======================= Generic cover fill from Excel =======================

def fill_cover_from_excel_generic(template_docx_path: str, excel_path_or_name: str, report_id: str, template_key: str) -> BytesIO:
    if not os.path.exists(template_docx_path):
        raise FileNotFoundError(f"Template not found: {template_docx_path}")

    df = _load_excel_df(excel_path_or_name)

    key_col = None
    for name in ["Report #", "Report#", "Report No", "Report no", "Report", "Report_No", "Report_Number"]:
        if name in df.columns:
            key_col = name
            break
    if not key_col:
        for c in df.columns:
            if "report" in c.lower():
                key_col = c
                break
    if not key_col:
        raise KeyError("Missing 'Report #' column (or variants) in Excel.")

    row_df = df.loc[df[key_col].astype(str).str.strip() == str(report_id or "").strip()]
    if row_df.empty:
        raise ValueError(f"Report # not found: {report_id}")
    row = row_df.iloc[0]
    excel_cols = list(df.columns)

    preferred = {
        "Result:": ["rating", "Rating", "RATING"],
        "Sample Description:": ["Item name / Description", "Item name/Description", "Description", "Item name", "Item Description"],
        "Item/ Material code:": ["Item#", "Item #", "Item code", "Item / Material code", "Item/ Material code"],
        "Category:": ["Category / Component name / Position ", "Category", "Component name", "Position"],
        "Collection:": ["Collection"],
        "Country of Destination:": ["Country of destination", "Country of Destination", "Destination"],
        "Supplier/ Subcontractor:": ["QA comment", "QA Comment", "QA comments", "Supplier / Subcontractor ", "Supplier/ Subcontractor", "Supplier", "Subcontractor"],
        "Customer:": ["Customer / Buyer", "Customer", "Buyer"],
        "Sample Size:": ["Sample Size", "Size"],
        "Sample Weight:": ["Sample Weight", "Weight"],
    }

    def _colnorm(x):
        s = unicodedata.normalize("NFD", str(x or "").strip().lower())
        s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
        s = re.sub(r"[^a-z0-9]+", " ", s)
        return re.sub(r"\s+", " ", s).strip()

    def _tokens(x):
        return set(_colnorm(x).split()) if x else set()

    def _overlap_score(label, colname):
        ln = _colnorm(label)
        cn = _colnorm(colname)
        if not ln or not cn:
            return 0.0
        score = 0.0
        if ln == cn:
            score += 5.0
        if ln in cn:
            score += 3.0
        if cn in ln:
            score += 1.0
        lt = _tokens(label)
        ct = _tokens(colname)
        if lt and ct:
            inter = lt & ct
            score += 4.0 * (len(inter) / max(1, len(lt))) + 2.0 * (len(inter) / max(1, len(ct)))
        return score

    def _pick_best_column(excel_columns, label_text: str, preferred_aliases=None) -> str:
        preferred_aliases = preferred_aliases or []
        is_item_code = _colnorm(label_text) in {"item material code", "item code", "item material"}
        blacklist = {"qr code"} if is_item_code else set()

        best_col, best_score = "", -1.0
        for c in excel_columns:
            if _colnorm(c) in blacklist:
                continue
            sc = _overlap_score(label_text, c)
            if c in preferred_aliases:
                sc += 2.0
            if sc > best_score:
                best_col, best_score = c, sc
        return best_col if best_score > 0.8 else ""

    def _val(row, colname: str) -> str:
        if not colname:
            return ""
        try:
            v = row.get(colname, "")
        except Exception:
            v = ""
        if pd is not None and hasattr(pd, "isna") and pd.isna(v):
            return ""
        return str(v).strip()

    doc = Document(template_docx_path)

    col_rating = next((c for c in excel_cols if _colnorm(c) == "rating"), "")
    if col_rating:
        _set_result_value(doc, _val(row, col_rating))

    tbl = _find_result_table(doc)
    if tbl is None:
        raise RuntimeError("Cover/RESULT table not found in template.")

    for r in tbl.rows:
        cells = r.cells
        if len(cells) >= 2:
            l_label = (cells[0].text or "").strip()
            if l_label.endswith(":"):
                cand = _pick_best_column(excel_cols, l_label, preferred_aliases=preferred.get(l_label))
                if cand:
                    cur = (cells[1].text or "").strip()
                    if _is_placeholder_dash(cur):
                        _set_cell_text_with_style(cells[1], cells[0], _val(row, cand))
        if len(cells) >= 4:
            r_label = (cells[2].text or "").strip()
            if r_label.endswith(":"):
                cand = _pick_best_column(excel_cols, r_label, preferred_aliases=preferred.get(r_label))
                if cand:
                    cur = (cells[3].text or "").strip()
                    if _is_placeholder_dash(cur):
                        _set_cell_text_with_style(cells[3], cells[2], _val(row, cand))

    _insert_overview_images_into_sample_picture(doc, report_id)

    _update_exec_summary_results_from_status(doc, report_id, template_key)
    _update_detail_results_and_comments(doc, report_id, template_key)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ======================= Wrappers / Entrypoints =======================

def fill_bed_cover_from_excel(template_docx_path: str, excel_path_or_name: str, report_id: str) -> BytesIO:
    return fill_cover_from_excel_generic(
        template_docx_path=template_docx_path,
        excel_path_or_name=excel_path_or_name,
        report_id=report_id,
        template_key="bed",
    )

def create_report_for_type(report_id: str, template_key: str, excel_path_or_name=None) -> BytesIO:
    """
    Locate template by TEMPLATE_MAP[template_key] next to app.py, then fill.
    Excel path is flexible: file or directory (auto-pick latest).
    """
    template_name = TEMPLATE_MAP.get(template_key)
    if not template_name:
        raise KeyError(f"Template not found for key: {template_key}")
    template_path = os.path.join(os.path.dirname(__file__), template_name)
    return fill_cover_from_excel_generic(
        template_docx_path=template_path,
        excel_path_or_name=excel_path_or_name or local_main,
        report_id=report_id,
        template_key=template_key,
    )
# ========================= END OF FILE =========================